from typing import Dict, Any
from .base import FileProcessor
import os
from docx import Document

class DOCXProcessor(FileProcessor):
    """DOCX文件处理器"""
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        处理DOCX文件
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            Dict[str, Any]: 包含文本内容和元数据的字典
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # 提取段落文本
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 提取表格数据
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            metadata = {
                'file_type': 'docx',
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'content_length': len(text)
            }
            
            return {
                'text': text,
                'tables': tables_content,
                'metadata': metadata,
                'success': True
            }
        except Exception as e:
            return {
                'text': '',
                'tables': [],
                'metadata': {},
                'success': False,
                'error': str(e)
            }
