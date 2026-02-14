#!/usr/bin/env python3
"""
文档加载器
用于加载不同格式的文档文件（Word、Excel等）
"""

from typing import Dict, Any, List
import os
import logging
from pathlib import Path

# 导入文档处理库
try:
    import docx
    from docx import Document as DocxDocument
except ImportError:
    logging.warning("python-docx not installed, Word document support may be limited")
    DocxDocument = None

try:
    import pandas as pd
except ImportError:
    logging.warning("pandas not installed, Excel document support may be limited")
    pd = None

logger = logging.getLogger(__name__)

class DocumentLoader:
    """文档加载器类，支持多种文档格式"""
    
    def __init__(self):
        """初始化文档加载器"""
        self.supported_formats = {
            ".docx": self._load_docx,
            ".xlsx": self._load_xlsx,
            ".txt": self._load_txt
        }
        logger.info("DocumentLoader initialized with supported formats: %s", list(self.supported_formats.keys()))
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        加载文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容和元数据的字典
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"不支持的文件格式: {file_extension}")
            
            logger.info("Loading document: %s", file_path)
            
            # 调用对应格式的加载函数
            load_func = self.supported_formats[file_extension]
            document_data = load_func(file_path)
            
            # 添加元数据
            document_data["file_path"] = str(file_path)
            document_data["file_name"] = file_path.name
            document_data["file_format"] = file_extension
            
            logger.info("Document loaded successfully: %s", file_path.name)
            return document_data
            
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {str(e)}")
            raise
    
    def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        加载Word文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容的字典
        """
        if DocxDocument is None:
            raise ImportError("python-docx not installed, cannot load Word documents")
        
        doc = DocxDocument(file_path)
        content = []
        
        # 读取段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # 读取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "content": "\n\n".join(content),
            "tables": tables,
            "type": "word_document"
        }
    
    def _load_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """
        加载Excel文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容的字典
        """
        if pd is None:
            raise ImportError("pandas not installed, cannot load Excel documents")
        
        # 读取所有工作表
        excel_data = pd.ExcelFile(file_path)
        sheets = {}
        
        for sheet_name in excel_data.sheet_names:
            # 读取工作表
            df = pd.read_excel(excel_data, sheet_name)
            
            # 转换为字典格式
            sheet_data = {
                "data": df.to_dict(orient="records"),
                "columns": list(df.columns),
                "shape": df.shape
            }
            sheets[sheet_name] = sheet_data
        
        # 提取文本内容（如果有）
        content = []
        for sheet_name, sheet_data in sheets.items():
            # 将表格数据转换为文本描述
            content.append(f"=== 工作表: {sheet_name} ===")
            content.append(f"列: {', '.join(sheet_data['columns'])}")
            content.append(f"行数: {sheet_data['shape'][0]}, 列数: {sheet_data['shape'][1]}")
            
            # 只提取前几行作为内容预览
            preview_rows = min(10, sheet_data['shape'][0])
            for i in range(preview_rows):
                row = sheet_data['data'][i]
                row_text = ", ".join([f"{col}: {row.get(col)}" for col in sheet_data['columns']])
                content.append(f"行 {i+1}: {row_text}")
            
            if sheet_data['shape'][0] > preview_rows:
                content.append(f"... 还有 {sheet_data['shape'][0] - preview_rows} 行数据")
            
            content.append("")
        
        return {
            "content": "\n".join(content),
            "sheets": sheets,
            "type": "excel_document"
        }
    
    def _load_txt(self, file_path: Path) -> Dict[str, Any]:
        """
        加载文本文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容的字典
        """
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return {
            "content": content,
            "type": "text_document"
        }
    
    def load_multiple_documents(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        加载目录中的多个文档
        
        Args:
            directory_path: 目录路径
            
        Returns:
            List[Dict]: 包含多个文档数据的列表
        """
        documents = []
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"目录不存在或不是目录: {directory_path}")
        
        # 遍历目录中的所有文件
        for file_path in directory_path.iterdir():
            if file_path.is_file():
                try:
                    file_extension = file_path.suffix.lower()
                    if file_extension in self.supported_formats:
                        document = self.load_document(file_path)
                        documents.append(document)
                        logger.info("Loaded document: %s", file_path.name)
                    else:
                        logger.info("Skipping unsupported file format: %s", file_path.name)
                except Exception as e:
                    logger.error(f"Failed to load document {file_path.name}: {str(e)}")
                    continue
        
        return documents

# 创建单例实例
document_loader = DocumentLoader()