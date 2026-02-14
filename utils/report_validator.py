#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告验证模块
用于验证生成的报告的质量和正确性
"""

import os
import re
from typing import List, Dict, Any
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd


class ReportValidator:
    """
    报告验证器，用于验证生成的报告的质量和正确性
    """
    
    def __init__(self):
        """
        初始化报告验证器
        """
        self.validation_results = []
    
    def validate(self, report_content: str, core_topics: List[str] = None, required_terms: List[str] = None) -> Dict[str, Any]:
        """
        验证报告内容的质量和正确性
        
        Args:
            report_content: 报告内容
            core_topics: 核心主题列表（可选）
            required_terms: 必需包含的术语列表（可选）
            
        Returns:
            Dict: 验证结果
        """
        try:
            # 基本验证
            validation_result = {
                "is_valid": True,
                "word_count": len(report_content),
                "sections": [],
                "coverage": {
                    "geopolitical": False,
                    "economic": False,
                    "market_sentiment": False
                },
                "errors": []
            }
            
            # 检查关键字覆盖
            if "地缘政治" in report_content:
                validation_result["coverage"]["geopolitical"] = True
            if "经济" in report_content or "数据" in report_content:
                validation_result["coverage"]["economic"] = True
            if "市场情绪" in report_content or "投资者" in report_content:
                validation_result["coverage"]["market_sentiment"] = True
            
            # 检查报告长度
            if len(report_content) < 1000:
                validation_result["errors"].append("报告长度过短")
                validation_result["is_valid"] = False
            
            # 检查核心主题
            if core_topics:
                for topic in core_topics:
                    if topic not in report_content:
                        validation_result["errors"].append(f"缺少核心主题：{topic}")
                        validation_result["is_valid"] = False
            
            # 检查必需术语
            if required_terms:
                for term in required_terms:
                    if term not in report_content:
                        validation_result["errors"].append(f"缺少必需术语：{term}")
                        validation_result["is_valid"] = False
            
            return validation_result
            
        except Exception as e:
            return {
                "is_valid": False,
                "word_count": 0,
                "sections": [],
                "coverage": {
                    "geopolitical": False,
                    "economic": False,
                    "market_sentiment": False
                },
                "errors": [f"验证失败：{str(e)}"]
            }
    
    def validate_directory(self, doc_path: str) -> Dict[str, Any]:
        """
        验证目录功能
        
        Args:
            doc_path: Word文档路径
            
        Returns:
            Dict: 验证结果
        """
        try:
            doc = Document(doc_path)
            
            # 查找TOC字段
            has_toc = False
            toc_paragraphs = []
            
            for para in doc.paragraphs:
                # 检查段落中是否包含"目录"或"Table of Contents"标题
                if '目录' in para.text or 'Table of Contents' in para.text:
                    toc_paragraphs.append(para)
                
                for run in para.runs:
                    # 检查是否包含TOC字段
                    try:
                        for element in run._r:
                            # 查找所有fldChar元素
                            if element.tag == qn('w:fldChar'):
                                # 查找段落中的所有instrText元素
                                for instr_element in run._r:
                                    if instr_element.tag == qn('w:instrText'):
                                        if 'TOC' in instr_element.text:
                                            has_toc = True
                                            toc_paragraphs.append(para)
                                            break
                            if has_toc:
                                break
                    except Exception as e:
                        continue
                if has_toc:
                    break
            
            # 计算实际标题数量（使用Heading样式的段落）
            actual_headings = []
            heading_counts = {1: 0, 2: 0, 3: 0}
            
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    heading_level = int(para.style.name.replace('Heading ', ''))
                    if 1 <= heading_level <= 3:
                        heading_counts[heading_level] += 1
                        actual_headings.append(para)
            
            total_headings = sum(heading_counts.values())
            
            # 检查目录项数量与标题数量是否匹配
            # 注意：实际目录项数量需要在Word中打开时才会更新，这里只检查结构
            
            return {
                'has_toc': has_toc,
                'has_toc_field': len(toc_paragraphs) > 0,
                'heading_counts': heading_counts,
                'total_headings': total_headings,
                'status': 'PASS' if has_toc and total_headings > 0 else 'FAIL',
                'message': '目录功能验证通过' if has_toc else '缺少TOC字段',
                'details': {
                    'toc_paragraph_count': len(toc_paragraphs),
                    'heading_1_count': heading_counts[1],
                    'heading_2_count': heading_counts[2],
                    'heading_3_count': heading_counts[3]
                }
            }
        
        except Exception as e:
            return {
                'has_toc': False,
                'has_toc_field': False,
                'heading_counts': {1: 0, 2: 0, 3: 0},
                'total_headings': 0,
                'status': 'ERROR',
                'message': f'验证失败：{str(e)}',
                'details': {}
            }
    
    def validate_data_accuracy(self, doc_path: str, source_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        验证数据准确性
        
        Args:
            doc_path: Word文档路径
            source_data: 源数据字典
            
        Returns:
            Dict: 验证结果
        """
        try:
            doc = Document(doc_path)
            
            # 提取文档中的所有文本
            doc_text = ""
            for para in doc.paragraphs:
                doc_text += para.text + "\n"
            
            # 验证关键指标
            verified_metrics = []
            failed_metrics = []
            
            # 如果有用户数据，验证用户数量
            if 'user_data' in source_data:
                user_df = source_data['user_data']
                if not user_df.empty:
                    # 验证总用户数
                    total_users = user_df['user_id'].nunique()
                    if str(total_users) in doc_text:
                        verified_metrics.append(f'总用户数: {total_users}')
                    else:
                        failed_metrics.append(f'总用户数: {total_users}')
            
            # 如果有每日数据，验证关键指标
            if 'daily_data' in source_data:
                daily_df = source_data['daily_data']
                if not daily_df.empty:
                    # 验证平均活跃用户数
                    avg_active_users = daily_df['active_users'].mean()
                    avg_active_users_rounded = round(avg_active_users)
                    if str(avg_active_users_rounded) in doc_text:
                        verified_metrics.append(f'平均活跃用户数: {avg_active_users_rounded}')
                    else:
                        failed_metrics.append(f'平均活跃用户数: {avg_active_users_rounded}')
            
            return {
                'status': 'PASS' if len(failed_metrics) == 0 else 'FAIL',
                'message': f'数据准确性验证通过' if len(failed_metrics) == 0 else f'有 {len(failed_metrics)} 个指标验证失败',
                'details': {
                    'verified_metrics': verified_metrics,
                    'failed_metrics': failed_metrics,
                    'total_verified': len(verified_metrics),
                    'total_failed': len(failed_metrics)
                }
            }
        
        except Exception as e:
            return {
                'status': 'ERROR',
                'message': f'验证失败：{str(e)}',
                'details': {}
            }
    
    def validate_format_standardization(self, doc_path: str) -> Dict[str, Any]:
        """
        验证格式标准化
        
        Args:
            doc_path: Word文档路径
            
        Returns:
            Dict: 验证结果
        """
        try:
            doc = Document(doc_path)
            
            # 检查页面设置
            has_proper_margins = True
            for section in doc.sections:
                # 检查页边距是否合理（1-3厘米）
                left_margin = section.left_margin.cm
                right_margin = section.right_margin.cm
                top_margin = section.top_margin.cm
                bottom_margin = section.bottom_margin.cm
                
                if not (1 <= left_margin <= 3 and 1 <= right_margin <= 3 and 
                        1 <= top_margin <= 3 and 1 <= bottom_margin <= 3):
                    has_proper_margins = False
                    break
            
            # 检查页眉页脚
            has_header = False
            has_footer = False
            
            for section in doc.sections:
                if section.header.paragraphs and section.header.paragraphs[0].text.strip():
                    has_header = True
                if section.footer.paragraphs and section.footer.paragraphs[0].text.strip():
                    has_footer = True
                break  # 只检查第一个节
            
            # 检查字体设置
            has_chinese_font = False
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.element.rPr and run.element.rPr.rFonts:
                        east_asia_font = run.element.rPr.rFonts.get(qn('w:eastAsia'))
                        if east_asia_font in ['宋体', '微软雅黑', 'SimSun', 'Microsoft YaHei']:
                            has_chinese_font = True
                            break
                if has_chinese_font:
                    break
            
            return {
                'status': 'PASS' if has_proper_margins and has_header and has_footer else 'FAIL',
                'message': '格式标准化验证通过' if has_proper_margins and has_header and has_footer else '格式不符合标准',
                'details': {
                    'has_proper_margins': has_proper_margins,
                    'has_header': has_header,
                    'has_footer': has_footer,
                    'has_chinese_font': has_chinese_font
                }
            }
        
        except Exception as e:
            return {
                'status': 'ERROR',
                'message': f'验证失败：{str(e)}',
                'details': {}
            }
    
    def validate_report(self, doc_path: str, source_data: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        全面验证报告
        
        Args:
            doc_path: Word文档路径
            source_data: 源数据字典（可选）
            
        Returns:
            Dict: 综合验证结果
        """
        validation_results = {
            'directory_validation': self.validate_directory(doc_path),
            'format_validation': self.validate_format_standardization(doc_path)
        }
        
        # 如果提供了源数据，进行数据准确性验证
        if source_data:
            validation_results['data_validation'] = self.validate_data_accuracy(doc_path, source_data)
        
        # 计算总体状态
        overall_status = 'PASS'
        for validation_type, result in validation_results.items():
            if result['status'] == 'FAIL':
                overall_status = 'FAIL'
                break
            elif result['status'] == 'ERROR':
                overall_status = 'ERROR'
                break
        
        return {
            'overall_status': overall_status,
            'validation_results': validation_results,
            'timestamp': pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')
        }
    
    def generate_validation_report(self, validation_result: Dict[str, Any], output_path: str = None) -> str:
        """
        生成验证报告
        
        Args:
            validation_result: 验证结果
            output_path: 输出路径（可选）
            
        Returns:
            str: 验证报告内容
        """
        report_content = "# 报告验证结果\n\n"
        report_content += f"验证时间: {validation_result['timestamp']}\n"
        report_content += f"总体状态: {validation_result['overall_status']}\n\n"
        
        for validation_type, result in validation_result['validation_results'].items():
            report_content += f"## {validation_type.replace('_', ' ').title()}\n"
            report_content += f"状态: {result['status']}\n"
            report_content += f"消息: {result['message']}\n"
            
            if 'details' in result:
                for detail_name, detail_value in result['details'].items():
                    if isinstance(detail_value, dict):
                        report_content += f"  - {detail_name}:\n"
                        for key, value in detail_value.items():
                            report_content += f"    - {key}: {value}\n"
                    elif isinstance(detail_value, list):
                        report_content += f"  - {detail_name}:\n"
                        for item in detail_value:
                            report_content += f"    - {item}\n"
                    else:
                        report_content += f"  - {detail_name}: {detail_value}\n"
            report_content += "\n"
        
        # 如果提供了输出路径，保存报告
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
        
        return report_content


# 测试代码
if __name__ == "__main__":
    validator = ReportValidator()
    
    # 示例用法
    doc_path = "report_output/chinese_technical_report.docx"
    
    if os.path.exists(doc_path):
        # 创建示例源数据
        source_data = {
            'user_data': pd.DataFrame({
                'user_id': [1, 2, 3, 4, 5],
                'user_name': ['用户1', '用户2', '用户3', '用户4', '用户5']
            }),
            'daily_data': pd.DataFrame({
                'date': pd.date_range('2026-02-01', periods=7),
                'active_users': [100, 120, 110, 130, 140, 150, 160]
            })
        }
        
        # 验证报告
        validation_result = validator.validate_report(doc_path, source_data)
        
        # 生成验证报告
        report_content = validator.generate_validation_report(validation_result)
        print(report_content)
        
        # 保存验证报告
        validator.generate_validation_report(validation_result, "report_output/report_validation.txt")
    else:
        print(f"文档不存在: {doc_path}")
