#!/usr/bin/env python3
"""
增强的格式转换器模块
负责将报告内容转换为多种格式，支持高级功能
"""

from typing import Dict, Any, List, Optional
import logging
from datetime import datetime
import html
import markdown as md
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)


class EnhancedFormatConverter:
    """增强的格式转换器，负责将报告内容转换为多种格式"""
    
    def __init__(self):
        """初始化增强的格式转换器"""
        self.supported_formats = [
            'markdown', 'md',
            'html',
            'plain_text', 'txt',
            'docx', 'word',
            'json',
            'latex'
        ]
        
        # 默认样式配置
        self.style_configs = {
            'markdown': {
                'title_level': 1,
                'section_level': 2,
                'subsection_level': 3,
                'include_metadata': True
            },
            'html': {
                'template': 'default',
                'include_css': True,
                'include_js': False,
                'responsive': True,
                'theme': 'light'
            },
            'docx': {
                'page_size': 'A4',
                'margin_top': 1.0,
                'margin_bottom': 1.0,
                'margin_left': 1.25,
                'margin_right': 1.25,
                'font_size': 11,
                'font_name': 'Arial'
            }
        }
        
        logger.info("EnhancedFormatConverter初始化完成")
    
    def can_convert(self, format_type: str) -> bool:
        """
        检查是否支持该格式
        
        Args:
            format_type: 格式类型
            
        Returns:
            如果支持返回True，否则返回False
        """
        return format_type.lower() in self.supported_formats
    
    def convert(self, sections: List[Dict[str, Any]], format_type: str = 'markdown', metadata: Dict[str, Any] = None, style_config: Dict[str, Any] = None) -> Any:
        """
        将报告章节转换为指定格式
        
        Args:
            sections: 报告章节
            format_type: 目标格式
            metadata: 报告元数据
            style_config: 样式配置
            
        Returns:
            转换后的内容，根据格式类型可能是字符串或文件对象
        """
        if not self.can_convert(format_type):
            logger.error(f"不支持的格式: {format_type}")
            return None
        
        format_type = format_type.lower()
        
        # 合并样式配置
        if style_config:
            for format_key, config in style_config.items():
                if format_key in self.style_configs:
                    self.style_configs[format_key].update(config)
        
        if format_type in ['markdown', 'md']:
            return self._convert_to_markdown(sections, metadata, self.style_configs['markdown'])
        elif format_type == 'html':
            return self._convert_to_html(sections, metadata, self.style_configs['html'])
        elif format_type in ['plain_text', 'txt']:
            return self._convert_to_plain_text(sections, metadata)
        elif format_type in ['docx', 'word']:
            return self._convert_to_docx(sections, metadata, self.style_configs['docx'])
        elif format_type == 'json':
            return self._convert_to_json(sections, metadata)
        elif format_type == 'latex':
            return self._convert_to_latex(sections, metadata)
        else:
            logger.error(f"无法转换为格式: {format_type}")
            return None
    
    def _convert_to_markdown(self, sections: List[Dict[str, Any]], metadata: Dict[str, Any] = None, style_config: Dict[str, Any] = None) -> str:
        """
        转换为Markdown格式
        
        Args:
            sections: 报告章节
            metadata: 报告元数据
            style_config: 样式配置
            
        Returns:
            Markdown格式的报告
        """
        markdown_content = ""
        
        # 添加报告标题
        if metadata and 'title' in metadata:
            title_level = style_config.get('title_level', 1)
            markdown_content += f"{'#' * title_level} {metadata['title']}\n\n"
        
        # 添加报告元数据
        if metadata and style_config.get('include_metadata', True):
            markdown_content += "## 报告信息\n\n"
            for key, value in metadata.items():
                if key != 'title':  # 标题已经单独显示
                    markdown_content += f"- **{key}**: {value}\n"
            markdown_content += "\n"
        
        # 添加章节内容
        section_level = style_config.get('section_level', 2)
        subsection_level = style_config.get('subsection_level', 3)
        
        for section in sections:
            # 添加章节标题
            markdown_content += f"{'#' * section_level} {section['title']}\n\n"
            
            # 添加章节内容
            if 'content' in section and section['content']:
                content = section['content']
                
                # 确保内容以换行符结尾
                if not content.endswith('\n'):
                    content += '\n'
                
                markdown_content += content
                
                # 确保章节之间有两个空行
                if not markdown_content.endswith('\n\n'):
                    markdown_content += '\n'
        
        return markdown_content.strip()
    
    def _convert_to_html(self, sections: List[Dict[str, Any]], metadata: Dict[str, Any] = None, style_config: Dict[str, Any] = None) -> str:
        """
        转换为HTML格式
        
        Args:
            sections: 报告章节
            metadata: 报告元数据
            style_config: 样式配置
            
        Returns:
            HTML格式的报告
        """
        theme = style_config.get('theme', 'light')
        include_css = style_config.get('include_css', True)
        responsive = style_config.get('responsive', True)
        
        # 开始HTML文档
        html_content = '<!DOCTYPE html>\n<html lang="zh-CN">\n<head>\n'
        html_content += '<meta charset="UTF-8">\n'
        html_content += '<meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        
        # 添加标题
        if metadata and 'title' in metadata:
            html_content += f'<title>{html.escape(metadata["title"])}</title>\n'
        else:
            html_content += '<title>报告</title>\n'
        
        # 添加CSS样式
        if include_css:
            html_content += '<style>\n'
            
            # 基础样式
            html_content += 'body { font-family: Arial, sans-serif; line-height: 1.6; margin: 0; padding: 0; }\n'
            
            # 根据主题添加样式
            if theme == 'dark':
                html_content += 'body { background-color: #121212; color: #e0e0e0; }\n'
                html_content += 'header { background-color: #1e1e1e; border-bottom: 1px solid #333; }\n'
                html_content += '.report-container { background-color: #1e1e1e; }\n'
                html_content += 'section { border-bottom: 1px solid #333; }\n'
            else:  # light theme
                html_content += 'body { background-color: #f5f5f5; color: #333; }\n'
                html_content += 'header { background-color: #fff; border-bottom: 1px solid #eee; }\n'
                html_content += '.report-container { background-color: #fff; }\n'
                html_content += 'section { border-bottom: 1px solid #eee; }\n'
            
            # 容器样式
            html_content += '.report-container { max-width: 900px; margin: 0 auto; padding: 20px; }\n'
            
            # 标题样式
            html_content += 'header h1 { margin: 0; padding: 20px 0; font-size: 2em; }\n'
            
            # 元数据样式
            html_content += '.metadata { margin-bottom: 30px; padding: 15px; border-radius: 5px; background-color: rgba(0, 0, 0, 0.05); }\n'
            html_content += '.metadata p { margin: 5px 0; }\n'
            
            # 章节样式
            html_content += 'section { padding: 20px 0; }\n'
            html_content += 'section:last-child { border-bottom: none; }\n'
            html_content += 'section h2 { margin-top: 0; font-size: 1.5em; }\n'
            html_content += 'section h3 { font-size: 1.2em; }\n'
            
            # 响应式设计
            if responsive:
                html_content += '@media (max-width: 768px) {\n'
                html_content += '  .report-container { padding: 10px; }\n'
                html_content += '  header h1 { font-size: 1.8em; }\n'
                html_content += '  section h2 { font-size: 1.3em; }\n'
                html_content += '  section h3 { font-size: 1.1em; }\n'
                html_content += '}\n'
            
            html_content += '</style>\n'
        
        html_content += '</head>\n<body>\n'
        
        # 添加报告容器
        html_content += '<div class="report-container">\n'
        
        # 添加报告标题
        if metadata and 'title' in metadata:
            html_content += '<header>\n'
            html_content += f'<h1>{html.escape(metadata["title"])}</h1>\n'
            html_content += '</header>\n'
        
        # 添加报告元数据
        if metadata:
            html_content += '<div class="metadata">\n'
            for key, value in metadata.items():
                if key != 'title':  # 标题已经单独显示
                    html_content += f'<p><strong>{html.escape(str(key))}:</strong> {html.escape(str(value))}</p>\n'
            html_content += '</div>\n'
        
        # 添加章节内容
        for section in sections:
            html_content += '<section>\n'
            html_content += f'<h2>{html.escape(section["title"])}</h2>\n'
            
            if 'content' in section and section['content']:
                # 转换Markdown内容为HTML
                content_html = md.markdown(section['content']) if 'markdown' in section.get('format', '') else section['content']
                html_content += f'<div class="content">{content_html}</div>\n'
            
            html_content += '</section>\n'
        
        html_content += '</div>\n'
        html_content += '</body>\n</html>'
        
        return html_content
    
    def _convert_to_plain_text(self, sections: List[Dict[str, Any]], metadata: Dict[str, Any] = None) -> str:
        """
        转换为纯文本格式
        
        Args:
            sections: 报告章节
            metadata: 报告元数据
            
        Returns:
            纯文本格式的报告
        """
        text_content = ""
        
        # 添加报告标题
        if metadata and 'title' in metadata:
            text_content += f"{'='*60}\n"
            text_content += f"{metadata['title']}\n"
            text_content += f"{'='*60}\n\n"
        
        # 添加报告元数据
        if metadata:
            text_content += f"{'_'*40}\n"
            text_content += f"报告信息\n"
            text_content += f"{'_'*40}\n"
            for key, value in metadata.items():
                if key != 'title':  # 标题已经单独显示
                    text_content += f"{key}: {value}\n"
            text_content += "\n"
        
        # 添加章节内容
        for section in sections:
            text_content += f"{'='*50}\n"
            text_content += f"{section['title']}\n"
            text_content += f"{'='*50}\n\n"
            
            if 'content' in section and section['content']:
                # 从HTML或Markdown中提取纯文本
                content = self._extract_plain_text(section['content'])
                text_content += f"{content}\n\n"
        
        return text_content.strip()
    
    def _convert_to_docx(self, sections: List[Dict[str, Any]], metadata: Dict[str, Any] = None, style_config: Dict[str, Any] = None) -> Document:
        """
        转换为Word文档格式
        
        Args:
            sections: 报告章节
            metadata: 报告元数据
            style_config: 样式配置
            
        Returns:
            Document对象
        """
        try:
            # 创建新文档
            doc = Document()
            
            # 设置页面大小和边距
            section = doc.sections[0]
            
            # 设置页边距（英寸）
            margin_top = style_config.get('margin_top', 1.0)
            margin_bottom = style_config.get('margin_bottom', 1.0)
            margin_left = style_config.get('margin_left', 1.25)
            margin_right = style_config.get('margin_right', 1.25)
            
            section.top_margin = Inches(margin_top)
            section.bottom_margin = Inches(margin_bottom)
            section.left_margin = Inches(margin_left)
            section.right_margin = Inches(margin_right)
            
            # 设置字体
            font_size = style_config.get('font_size', 11)
            font_name = style_config.get('font_name', 'Arial')
            
            # 设置正文样式
            style = doc.styles['Normal']
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            
            # 设置段落样式
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = 1.15
            
            # 添加报告标题
            if metadata and 'title' in metadata:
                title_paragraph = doc.add_heading(metadata['title'], level=0)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加空行
                doc.add_paragraph()
            
            # 添加报告元数据
            if metadata:
                metadata_paragraph = doc.add_paragraph()
                metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for key, value in metadata.items():
                    if key != 'title':  # 标题已经单独显示
                        metadata_paragraph.add_run(f"{key}: {value}    ")
                
                # 添加分页符
                self._add_page_break(doc)
            
            # 添加章节内容
            for section in sections:
                # 添加章节标题
                heading_paragraph = doc.add_heading(section['title'], level=1)
                
                # 添加章节内容
                if 'content' in section and section['content']:
                    # 将Markdown内容转换为Word格式
                    self._add_markdown_content(doc, section['content'], font_name, font_size)
                
                # 添加分页符（除了最后一个章节）
                if section != sections[-1]:
                    self._add_page_break(doc)
            
            logger.info("Word文档转换完成")
            return doc
            
        except Exception as e:
            logger.error(f"转换为Word文档失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _convert_to_json(self, sections: List[Dict[str, Any]], metadata: Dict[str, Any] = None) -> str:
        """
        转换为JSON格式
        
        Args:
            sections: 报告章节
            metadata: 报告元数据
            
        Returns:
            JSON格式的报告
        """
        import json
        
        # 创建报告数据结构
        report_data = {
            'metadata': metadata or {},
            'sections': sections,
            'generated_at': datetime.now().isoformat()
        }
        
        # 转换为JSON字符串
        return json.dumps(report_data, ensure_ascii=False, indent=2)
    
    def _convert_to_latex(self, sections: List[Dict[str, Any]], metadata: Dict[str, Any] = None) -> str:
        """
        转换为LaTeX格式
        
        Args:
            sections: 报告章节
            metadata: 报告元数据
            
        Returns:
            LaTeX格式的报告
        """
        latex_content = "\\documentclass{article}\n\n"
        latex_content += "\\usepackage[UTF8]{ctex}\n"
        latex_content += "\\usepackage{a4wide}\n"
        latex_content += "\\usepackage{amsmath}\n"
        latex_content += "\\usepackage{graphicx}\n"
        latex_content += "\\usepackage{booktabs}\n"
        latex_content += "\\usepackage{listings}\n\n"
        
        # 添加文档配置
        latex_content += "\title{"
        if metadata and 'title' in metadata:
            latex_content += html.escape(metadata['title'])
        else:
            latex_content += "报告"
        latex_content += "}\\n"
        
        latex_content += "\\author{"
        if metadata and 'author' in metadata:
            latex_content += html.escape(metadata['author'])
        latex_content += "}\\n"
        
        latex_content += "\\date{"
        if metadata and 'date' in metadata:
            latex_content += html.escape(metadata['date'])
        else:
            latex_content += "\\today"
        latex_content += "}\\n\n"
        
        latex_content += "\\begin{document}\\n\\maketitle\\n"
        
        # 添加报告元数据
        if metadata:
            latex_content += "\\section*{报告信息}\\n"
            latex_content += "\\begin{itemize}\\n"
            for key, value in metadata.items():
                if key not in ['title', 'author', 'date']:  # 这些已经单独显示
                    latex_content += f"  \\item \\textbf{{{html.escape(str(key))}}}: {html.escape(str(value))}\\n"
            latex_content += "\\end{itemize}\\n\\n"
        
        # 添加章节内容
        for i, section in enumerate(sections):
            latex_content += f"\\section{{{html.escape(section['title'])}}}\\n"
            
            if 'content' in section and section['content']:
                # 简单的Markdown到LaTeX转换
                content = section['content']
                
                # 替换标题
                content = content.replace('# ', '\\subsection*{')
                content = content.replace('#', '')
                content = content.replace('\n', '\\n')
                
                # 替换粗体
                content = content.replace('**', '\\textbf{')
                content = content.replace('**', '}')
                
                # 替换斜体
                content = content.replace('*', '\\textit{')
                content = content.replace('*', '}')
                
                # 添加内容
                latex_content += f"{content}\\n\\n"
        
        latex_content += "\\end{document}"
        
        return latex_content
    
    def _extract_plain_text(self, content: str) -> str:
        """
        从HTML或Markdown中提取纯文本
        
        Args:
            content: 包含HTML或Markdown的内容
            
        Returns:
            纯文本
        """
        import re
        
        # 移除HTML标签
        text = re.sub(r'<[^>]+>', '', content)
        
        # 移除Markdown格式
        text = re.sub(r'#+', '', text)  # 标题
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # 粗体
        text = re.sub(r'\*([^*]+)\*', r'\1', text)  # 斜体
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # 链接
        text = re.sub(r'!\[([^\]]+)\]\([^)]+\)', r'\1', text)  # 图片
        text = re.sub(r'-\s*', '', text)  # 列表项
        text = re.sub(r'\d+\.\s*', '', text)  # 编号列表
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # 代码块
        text = re.sub(r'`([^`]+)`', r'\1', text)  # 行内代码
        
        return text
    
    def _add_markdown_content(self, doc: Document, content: str, font_name: str, font_size: int):
        """
        将Markdown内容添加到Word文档
        
        Args:
            doc: Document对象
            content: Markdown内容
            font_name: 字体名称
            font_size: 字体大小
        """
        import re
        
        # 分割内容为段落
        paragraphs = re.split(r'\n\s*\n', content)
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # 检查是否是标题
            heading_match = re.match(r'(#+)\s*(.*)', para)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                elif level == 3:
                    doc.add_heading(text, level=3)
                else:
                    doc.add_heading(text, level=4)
            
            # 检查是否是列表
            elif re.match(r'^(-|\d+\.)\s+', para):
                # 简单处理列表
                items = re.split(r'\n(?=-|\d+\.)', para)
                for item in items:
                    # 移除列表标记
                    item_text = re.sub(r'^(-|\d+\.)\s+', '', item)
                    
                    # 添加列表项
                    doc.add_paragraph(item_text, style='List Bullet' if item.startswith('-') else 'List Number')
            
            # 检查是否是代码块
            elif para.startswith('```'):
                # 简单处理代码块
                code_text = re.sub(r'^```\w*\n|```$', '', para)
                doc.add_paragraph(code_text, style='Code')
            
            # 普通段落
            else:
                # 处理内联格式
                para_text = para
                
                # 处理粗体
                para_text = re.sub(r'\*\*(.*?)\*\*', r'\1', para_text)
                
                # 处理斜体
                para_text = re.sub(r'\*(.*?)\*', r'\1', para_text)
                
                # 处理链接
                para_text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', para_text)
                
                # 处理图片
                para_text = re.sub(r'!\[(.*?)\]\((.*?)\)', r'\1', para_text)
                
                # 添加段落
                doc.add_paragraph(para_text)
    
    def _add_page_break(self, doc: Document):
        """
        在Word文档中添加分页符
        
        Args:
            doc: Document对象
        """
        run = doc.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def save_docx(self, doc: Document, file_path: str):
        """
        保存Word文档
        
        Args:
            doc: Document对象
            file_path: 文件路径
        """
        try:
            doc.save(file_path)
            logger.info(f"Word文档已保存到: {file_path}")
            return True
        except Exception as e:
            logger.error(f"保存Word文档失败: {e}")
            return False
    
    def set_style_config(self, format_type: str, config: Dict[str, Any]):
        """
        设置格式的样式配置
        
        Args:
            format_type: 格式类型
            config: 样式配置
        """
        format_type = format_type.lower()
        
        if format_type not in self.style_configs:
            self.style_configs[format_type] = {}
        
        self.style_configs[format_type].update(config)
        logger.info(f"已更新 {format_type} 格式的样式配置")
    
    def get_style_config(self, format_type: str) -> Optional[Dict[str, Any]]:
        """
        获取格式的样式配置
        
        Args:
            format_type: 格式类型
            
        Returns:
            样式配置，如果不存在返回None
        """
        return self.style_configs.get(format_type.lower())
    
    def list_supported_formats(self) -> List[str]:
        """
        列出所有支持的格式
        
        Returns:
            支持的格式列表
        """
        return self.supported_formats
